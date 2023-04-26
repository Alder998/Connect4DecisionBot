import docx
import pandas as pd
from fpdf import FPDF

def AutoCoverLetter ():

    ruolo = input('Inserisci Ruolo: ')
    azienda = input('Inserisci Azienda: ')

    doc = docx.Document(
        r"C:\Users\39328\OneDrive\Desktop\Davide\Cose Utili\Application 2021-2022\Davide_Ontano_Cover_Letter_2022_2023_per_algoritmo.docx")

    totalText = list()
    for paragraph in range(len(doc.paragraphs)):
        textEachParagraph = doc.paragraphs[paragraph].text
        totalText.append(textEachParagraph)

    totalText[3] = 'Object: ' + ruolo + ' - ' + azienda
    totalText[
        12] = 'I am sure that collaborating with ' + azienda + ' would certainly be an enriching opportunity to be a part of this challenging contest, in which I could easily adapt and where I could get the great chance to improve my professional and personal competences. '

    # Mettiamo tutto insieme
    # totalText = '\n'.join(totalText)

    print('\n')

    print('File Modificato')

    pdf = FPDF()

    # Add a page
    pdf.add_page()

    # set style and size of font
    # that you want in the pdf

    pdf.set_font("Helvetica", size=10)

    # add another cell

    for paragraph in range(len(totalText)):
        totalText[paragraph] = totalText[paragraph].replace(u"\u2018", "'").replace(u"\u2019", "'").replace(u"\u201c",
                                                                                                            "'").replace(
            u"\u201d", "'")
        pdf.multi_cell(190, 7, txt=totalText[paragraph], align='L')

    pdf.set_right_margin(20)
    pdf.set_right_margin(20)

    # save the pdf with name .pdf

    nomeFile = ruolo + ' - ' + azienda

    path = r"C:\Users\39328\OneDrive\Desktop\Davide\Cose Utili\Application 2021-2022\AutoCoverLetter\Cover Letter - " + nomeFile + '.pdf'

    pdf.output(path)

    print('\n')
    print('File Salvato')



